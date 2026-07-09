"""Builds the Word (.docx) file from the requirements data."""
from io import BytesIO
from docx import Document


def build_docx(data):
    """Turn the requirements data into a Word document (in memory)."""
    doc = Document()
    doc.add_heading("Requirements Document", 0)

    doc.add_heading("1. Requirements", level=1)
    doc.add_heading("Functional", level=2)
    for r in data["requirements"]["functional"]:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_heading("Non-functional", level=2)
    for r in data["requirements"]["non_functional"]:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("2. User Stories", level=1)
    for s in data["user_stories"]:
        doc.add_heading(f'{s["id"]} - {s["role"]}', level=3)
        doc.add_paragraph(s["story"])

    doc.add_heading("3. Acceptance Criteria", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Story", "Scenario", "Given", "When", "Then"]):
        table.rows[0].cells[i].text = h
    for ac in data["acceptance_criteria"]:
        cells = table.add_row().cells
        cells[0].text = str(ac.get("story_id", ""))
        cells[1].text = str(ac.get("scenario", ""))
        cells[2].text = str(ac.get("given", ""))
        cells[3].text = str(ac.get("when", ""))
        cells[4].text = str(ac.get("then", ""))

    doc.add_heading("4. Use Cases", level=1)
    uc_headers = ["ID", "Actors", "Description", "Preconditions",
                  "Trigger", "Main flow", "Alternative flow"]
    uc_keys = ["use_case_id", "actors", "description", "preconditions",
               "trigger", "main_flow", "alternative_flow"]
    table = doc.add_table(rows=1, cols=len(uc_headers))
    table.style = "Table Grid"
    for i, h in enumerate(uc_headers):
        table.rows[0].cells[i].text = h
    for uc in data.get("use_cases", []):
        cells = table.add_row().cells
        for i, key in enumerate(uc_keys):
            cells[i].text = str(uc.get(key, ""))

    doc.add_heading("5. Assumptions & Added Info", level=1)
    for a in data.get("assumptions", []):
        doc.add_paragraph(a, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf